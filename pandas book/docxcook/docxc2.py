from docx import Document
from docx.shared import Cm

'''
使用文本
块级与内联文本对象
段落Paragraph是Word中的主要块级对象

块级对象在它所包含的文本的左右边界之间流动，每当文本超出右边界时就添加一行

对于段落来说，边界通常是页边距。若页面是列布局，边界是列边界，如果段落在表格单元格内，边界是单元格边界

表格也是块级对象

内联对象是块级对象内容的一部分。例如，出现在黑体的单词或全大写的句子。最常见的内联对象是run。块容器中的所有内容都在内联对象中。通常，一个段落包含一个或多个run，每个run包含段落的一部分文本

块级对象的属性指定位置，比如段落前后的缩进。

内联对象的属性通常指定内容，比如字体、字体大小、粗体和斜体。
'''


'''新建文档'''
document = Document()

'''插入段落'''
paragraph = document.add_paragraph('段落1')
prior_paragraph = paragraph.insert_paragraph_before('段落0')  # 前插

'''插入标题'''
document.add_heading('The REAL meaning of the universe')  # 默认为标题 1
document.add_heading('The role of dolphins', level=2)

'''插入分页符'''
document.add_page_break()

'''插入表格'''
table = document.add_table(rows=2, cols=2)  # 2行2列的表格
cell = table.cell(0, 0)  # 第1行第1列的单元格
cell.text = '姓名'
cell = table.cell(0, 1)  # 第1行第2列的单元格
cell.text = '学号'
row = table.rows[1]  # 第2行
row.cells[0].text = '甲'
row.cells[1].text = '2015012755'
row = table.add_row()  # 插入一行
row = table.rows[2]  # 第3行
row.cells[0].text = '乙'
row.cells[1].text = '2015012756'

for row in table.rows:  # 遍历表格
    for cell in row.cells:
        print(cell.text, end=' ')
    print()

table = document.add_table(1, 3, style='Light Shading Accent 1')
heading_cells = table.rows[0].cells
heading_cells[0].text = '姓名'
heading_cells[1].text = '语文'
heading_cells[2].text = '数学'
grades = (
    ('甲', 90, 80),
    ('乙', 80, 90),
    ('丙', 100, 100),
)
for name, chinese, math in grades:
    cells = table.add_row().cells
    cells[0].text = name
    cells[1].text = str(chinese)
    cells[2].text = str(math)

# table.style = 'LightShading-Accent1'

'''插入图片'''
document.add_picture('11.png')
document.add_picture('11.png', width=Cm(4.0))  # 宽度为4cm

'''段落样式'''
document.add_paragraph('无序列表', style='ListBullet')  # 无序列表
paragraph = document.add_paragraph('无序列表')
paragraph.style = 'List Bullet'

'''加粗斜体'''
paragraph = document.add_paragraph('正文')
run = paragraph.add_run('加粗')  # 分步
run.bold = True
paragraph.add_run('正文')

paragraph = document.add_paragraph('正文')
paragraph.add_run('加粗').bold = True  # 一步到位
paragraph.add_run('斜体').italic = True

'''字符样式'''
paragraph = document.add_paragraph('正文')
run = paragraph.add_run('强调')  # 分步
run.style = 'Emphasis'

paragraph = document.add_paragraph('正文')
paragraph.add_run('强调', style='Emphasis')  # 一步到位

document.save('docx2.docx')  # 保存

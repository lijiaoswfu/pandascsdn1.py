from docx import Document
from docx.shared import Inches

'''
基本概念
主要对象

Document：Word文档对象
Paragraph：段落，一个回车一个段落
Run：节段，一个段落多个节段
'''
document = Document()

document.add_heading('Document Title', level=0)  # 插入标题

paragraph = document.add_paragraph('A plain paragraph having some ')  # 插入段落
paragraph.add_run('bold').bold = True  # 正文
paragraph.add_run(' and some ')
paragraph.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)  # 插入标题 1
document.add_paragraph('Intense quote', style='Intense Quote')  # 插入段落，明显引用
document.add_paragraph('first item in unordered list', style='List Bullet')  # 插入段落，无序列表
document.add_paragraph('first item in ordered list', style='List Number')  # 插入段落，有序列表

document.add_picture('11.png', width=Inches(5))  # 插入图片

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)  # 插入表格
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()  # 插入分页

document.save('docx1.docx')  # 保存

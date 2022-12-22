# coding:utf-8
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.shared import Pt
import pandas as pd
import numpy as np
import os
import xlwings as xw



app = xw.App(visible=False)
workbook = app.books.add()
worksheet = workbook.sheets.add('新工作表')

pd.set_option('display.max_columns', None)
pd.set_option('display.max_rows', None)



basicData = pd.read_csv('./Data/2023.csv',encoding='utf-8-sig')
shp = pd.read_excel('./Data/ceshi.xlsx',encoding='utf-8-sig')
shp.rename(columns = {"xinxuhao": "序号", "MIAN_JI":"面积"}, inplace=True)
shp['现勘时间'].astype('str')
mega = basicData[['序号','镇镇（街道','村（居委）','社（林班）','林地权属','地类','林种','起源','森林类别','公益林事权','国家公益林','优势树种','林地保护等','是否林地']]
shpMerge = pd.merge(shp,mega,how='left',on=['序号'])
# print(shpMerge.head())

shpfinall = shpMerge[['镇镇（街道','村（居委）','社（林班）','是否林地','细斑号','占地属性','面积','林地权属','地类','林种','起源','森林类别','公益林事权','国家公益林','优势树种','林地保护等','备注']]
shpfinall.reset_index(drop=True,inplace=True)
shpMergeunique = shpMerge['项目名称'].unique() # 获取项目不重复值
shpMergeuniquetime = shpMerge['现勘时间'].unique() # 获取时间不重复值
shpMergeuniquepepo = shpMerge['勘测人员'].unique() # 获取人员不重复值

ldwyzshp = shpMerge.query("是否林地 == '林地'")
ldwyzshpysp = ldwyzshp.query("占地属性 != '已审批'" )  # 条件筛选，林地并非已审批
ldwyzshpysp['WZ']= ldwyzshpysp['镇镇（街道']+ldwyzshpysp['村（居委）']+ldwyzshpysp['社（林班）'].map(str)
lzarry = ldwyzshpysp['林种'].unique() # 获取林种不重复值
dileiarry = ldwyzshpysp['地类'].unique() # 获取地类不重复值
ysszarry = ldwyzshpysp['优势树种'].unique() # 获取优势树种不重复值
sllbarry = ldwyzshpysp['森林类别'].unique() # 获取森林类别不重复值
wzarry = ldwyzshpysp['WZ'].unique() # 获取位置不重复值
bzarry = ldwyzshpysp['备注'].unique() # 获取位置不重复值

'''
for coluName in shpMergeunique:
    os.makedirs('./'+coluName, exist_ok=True)  # 建立对应的项目文件
'''

# shpMergegroupby = shpMergegroupby.rename(columns={'占地属性':'镇镇（街道'}) # 修改列名，以便让两个表列名贴合

# sMgsum = shpMergegroupby.to_frame()
# sMgsum['index']= range(1,len(sMgsum)+1)
# sMgsum.set_index('index',inplace = True)


#shpMergegroupby.loc["合计"] =shpMergegroupby.iloc[:,2].sum(axis=0)
# shpMergegroupby.loc['合计','占地属性'] = '合计'
# shpMergegroupby.loc['合计','地类'] = ''

sfld = shpfinall.groupby(['是否林地'],as_index=False)['面积'].sum()  # 根据是否林地统计面积,同时取消按条件索引
lmds = ''
j = 0
for i in range(len(sfld)):
   lmds = lmds + sfld.iloc[i,0] + sfld.iloc[i,1].astype(str) + '公顷' # 循环读取林地和非林地分类面积
   if j < len(sfld) -1:
      j=j+1
      lmds = lmds +"、"
hejisum = str(round(sfld.iloc[:,1].sum(),4))
sfld.loc["合计"] = sfld.iloc[:,1].sum(axis=0)
sfld.loc['合计','是否林地'] = '合计'





shpMergegroupby = shpfinall.groupby(['是否林地','占地属性','备注'],as_index=False)['面积'].sum()  # 根据是否林地统计面积

shequer = shpMergegroupby.query("是否林地 == '林地'")  # 条件索引为林地
shequysp = shequer.query("占地属性 != '已审批'" )  # 条件筛选，林地并非已审批



# 林分因子字段读取
beizhu = ''
j=0
for i in bzarry:
   beizhu = beizhu +str(i) # 循环读取备注属性
   if j < bzarry.size -1:
      j=j+1
      beizhu = beizhu +"、"

wztext = ''
j = 0
for i in wzarry:
   wztext = wztext +str(i)+"社" # 循环读取位置属性
   if j < wzarry.size -1:
      j=j+1
      wztext = wztext +"、"

sllbtext = ''
j = 0
for i in sllbarry:
   sllbtext = sllbtext +str(i) # 循环读取森林类别属性
   if j < sllbarry.size -1:
      j=j+1
      sllbtext = sllbtext +"、"

yssztext = ''
j = 0
for i in ysszarry:
   yssztext = yssztext +str(i) # 循环读取优势树种属性
   if j < ysszarry.size -1:
      j=j+1
      yssztext = yssztext +"、"

lztext = ''
j= 0
for i in lzarry:
   lztext = lztext +str(i) # 循环读取林种属性
   if j < lzarry.size -1:
      j=j+1
      lztext = lztext +"、"

dileitxt = ''
j = 0
for i in dileiarry:
   dileitxt = dileitxt +str(i) # 循环读取地类属性
   if j < dileiarry.size -1:
      j=j+1
      dileitxt = dileitxt +"、"

ldsx = ''
j = 0
for i in range(len(shequer)):
   ldsx = ldsx +shequer.iloc[i,1] +shequer.iloc[i,3].astype(str) + '公顷' # 循环读取林地占地属性分类面积
   if j < len(shequer) -1:
      j=j+1
      ldsx = ldsx +"、"

# shppovit = pd.pivot_table(shpfinall,index=['地类','占地属性'],values=['面积'],aggfunc=[np.sum],fill_value=0,margins=True) 数据透视表获取值
DF = pd.concat([shpfinall,sfld])
DF= DF[['镇镇（街道','村（居委）','社（林班）','是否林地','细斑号','占地属性','面积','林地权属','地类','林种','起源','森林类别','公益林事权','国家公益林','优势树种','林地保护等','备注']]
DF.reset_index(drop=True, inplace=True)
bgexcel = DF.shape[0]+3 #获取最终写入的行数
print(bgexcel)


# 开始EXCEL格式设置
worksheet.range('B1:R1').api.merge()
worksheet.range('B2:R2').api.merge()
worksheet.range('B1').value = shpMergeunique + '占地图班勘验一览表'
worksheet.range('B2').value = '单位：公顷、厘米'
worksheet.range('A3').value = DF
worksheet.api.Columns(1).Delete()
worksheet.range('A1').api.Font.Size = 20  #设置字号
worksheet.range('A1').api.Font.Bold = True  #设置粗体
worksheet.range('A1').api.HorizontalAlignment = -4108  # -4108 水平居中。 -4131 靠左，-4152 靠右。
worksheet.range('A2').api.HorizontalAlignment = -4152  # -4108 水平居中。 -4131 靠左，-4152 靠右。
worksheet.range('A1').api.VerticalAlignment = -4130      # -4108 垂直居中（默认）。 -4160 靠上，-4107 靠下， -4130 自动换行对齐。
Ryou = 'Q'+str(bgexcel)
print(Ryou)
bgexcelfw = worksheet.range('A3',Ryou)
bgexcelfw.api.HorizontalAlignment = -4108
bgexcelfw.api.VerticalAlignment = -4130

worksheet.autofit()
worksheet.range('A1').row_height =30
bgexcelfw.row_height = 30
"""设置边框"""
# Borders(9) 底部边框，LineStyle = 1 直线。
bgexcelfw.api.Borders(9).LineStyle = 1
bgexcelfw.api.Borders(9).Weight = 3                # 设置边框粗细。

# Borders(7) 左边框，LineStyle = 2 虚线。
bgexcelfw.api.Borders(7).LineStyle = 1
bgexcelfw.api.Borders(7).Weight = 3

# Borders(8) 顶部框，LineStyle = 5 双点划线。
bgexcelfw.api.Borders(8).LineStyle = 1
bgexcelfw.api.Borders(8).Weight = 3

# Borders(10) 右边框，LineStyle = 4 点划线。
bgexcelfw.api.Borders(10).LineStyle = 1
bgexcelfw.api.Borders(10).Weight = 3

# # Borders(11) 内部垂直边线。
bgexcelfw.api.Borders(11).LineStyle = 1
bgexcelfw.api.Borders(11).Weight = 2

# # Borders(12) 内部水平边线。
bgexcelfw.api.Borders(12).LineStyle = 1
bgexcelfw.api.Borders(12).Weight = 2

workbook.save('test.xlsx')
workbook.close()
app.quit()

# 开始WORD报告编辑
def docxhd(ph,st):
   runHead = ph.add_run(st)  # 变量代入内容
   runHead.font.name = '方正小标宋_GBK'
   runHead.font.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋_GBK')  # 设置中文字体
   runHead.font.size = Pt(22)  # 字体大小
   paragraph_formatHead = ph.paragraph_format
   paragraph_formatHead.line_spacing = Pt(30)  # 固定值，30磅
   ph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落居中对齐

def docxzw(pzw,szw):
   runFirst = pzw.add_run(szw)
   runFirst.font.name = '方正仿宋_GBK'
   runFirst.font.element.rPr.rFonts.set(qn('w:eastAsia'), '方正仿宋_GBK')
   runFirst.font.size = Pt(16)
   paragraph_formatFirst = pzw.paragraph_format
   paragraph_formatFirst.line_spacing = Pt(28)  # 固定值，30磅
   pzw.paragraph_format.first_line_indent = Cm(1.1)

def docxdw(pdw,sdw):
   runFirst = pdw.add_run(sdw)
   runFirst.font.name = '方正仿宋_GBK'
   runFirst.font.element.rPr.rFonts.set(qn('w:eastAsia'), '方正仿宋_GBK')
   runFirst.font.size = Pt(16)
   paragraph_formatFirst = pdw.paragraph_format
   paragraph_formatFirst.line_spacing = Pt(28)  # 固定值，30磅
   pdw.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 段落右对齐

doc = Document()
# print(doc.add_heading("一级标题", level=1))   添加一级标题的时候出错，还没有解决！
strHead = '重庆市涪陵区林业规划和资源监测中心关于'+ shpMergeunique +'占地的现场勘验报告'
paragraphHead = doc.add_paragraph()
docxhd(paragraphHead,strHead)


paragraphFirst = doc.add_paragraph()
strFirst = "按照区林业局安排，我中心就"+shpMergeunique+"占地一案，本着实事求是、客观公正的原则，严格按照相关调查技术规范，深入现场进行了实地勘验。现就勘验情况报告如下："
docxzw(paragraphFirst,strFirst)

p2 =doc.add_paragraph()
s2 ="一、勘验时间："+shpMergeuniquetime+"。"
docxzw(p2,s2)

p3 =doc.add_paragraph()
s3 = "二、勘验地点："+wztext+"。"
docxzw(p3,s3)

p4 =doc.add_paragraph()
s4 = "三、勘验人："+shpMergeuniquepepo+"。"
docxzw(p4,s4)

p5 = doc.add_paragraph()
s5 = "四、勘验内容：占用土地类型、地点、用途、面积，及林地林分调查因子。"
docxzw(p5,s5)

p6 = doc.add_paragraph()
s6 = "五、勘验方法：通过现场观察及了解访问，查阅比对卫片、航片及林业档案资料确定占用土地类型；现场核实占地地点、用途。对占地区域用无人机现场航拍实测，然后在1：1万地形图上对实拍场景进行勾绘区划，参照地形图上原有地类、行政区划界、2019年林地变更调查成果、森林分类经营资料、退耕还林等林业工程建设、自然保护区、审批边界等资料，以及占用地周边林地现状，综合分析确定占用林地的面积、林分因子。"
docxzw(p6,s6)

p7 = doc.add_paragraph()
s7 = "六、勘验结果："+shpMergeunique+"涉及面积"+hejisum+"公顷,其中"+str(lmds)+"。所涉及林地中"+ldsx+"。所占林地地类为"+dileitxt+"，林种为"+lztext+"，优势树种为"+yssztext+"，按森林类别分为"+sllbtext+"。所涉及林地现状为"+beizhu+"等。具体详见附表。"
docxzw(p7,s7)

p8 = doc.add_paragraph()
s8 = "附件：1."+shpMergeunique+"占用林地勘验图\n2. "+shpMergeunique+"占用林地小班一览表"
docxzw(p8,s8)

p9 = doc.add_paragraph()
s9 = "\n\n勘验人（签字）：\t\t\t\t"
docxdw(p9,s9)

p10 =doc.add_paragraph()
s10 = "重庆市涪陵区林业规划和资源监测中心"
docxdw(p10,s10)

p11 = doc.add_paragraph()
s11 = shpMergeuniquetime
docxdw(p11,s11)

doc.save("new1.docx")





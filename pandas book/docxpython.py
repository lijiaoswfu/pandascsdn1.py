# coding:utf-8
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.shared import Pt

doc = Document()
# print(doc.add_heading("一级标题", level=1))   添加一级标题的时候出错，还没有解决！
strHead = '重庆市涪陵区林业规划和资源监测中心关于曾淑明居民房占地的现场勘验报告'
paragraphHead = doc.add_paragraph()
runHead = paragraphHead.add_run(strHead) # 变量代入内容
runHead.font.name='方正小标宋_GBK'
runHead.font.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋_GBK')  # 设置中文字体
runHead.font.size = Pt(22) # 字体大小
paragraph_formatHead = paragraphHead.paragraph_format
paragraph_formatHead.line_spacing = Pt(30)  # 固定值，30磅
paragraphHead.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落居中对齐

paragraphFirst = doc.add_paragraph()
strFirst = "按照区林业局安排，我中心就曾淑明居民房占地一案，本着实事求是、客观公正的原则，严格按照相关调查技术规范，深入现场进行了实地勘验。现就勘验情况报告如下：\n一、勘验时间：2022年1月11日。\n二、勘验地点：大木乡武陵村1社。\n三、勘验人：李骄、周文杰。陪同人员：代朝伟（大木乡农服中心主任）\n四、勘验内容：占用土地类型、地点、用途、面积，及林地林分调查因子。\n五、勘验方法：通过现场观察及了解访问，查阅比对卫片、航片及林业档案资料确定占用土地类型；现场核实占地地点、用途。对占地区域用无人机现场航拍实测，然后在1：1万地形图上对实拍场景进行勾绘区划，参照地形图上原有地类、行政区划界、2019年林地变更调查成果、森林分类经营资料、退耕还林等林业工程建设、自然保护区、审批边界等资料，以及占用地周边林地现状，综合分析确定占用林地的面积、林分因子。\n六、勘验结果：曾淑明居民房涉及面积0.0442公顷，位于大木山自然保护区实验区内，涉及面积中0.0407公顷为非林地，0.0035公顷为林地。所占林地地类为国家特别规定灌木林地；林种为自然保护区林，总占林地的权属为集体；优势树种为其他灌木，起源为天然；按森林类别分为国家公益林；林地保护等级为Ⅱ。所占林地主要用途为建房。具体详见附表。\n\n附件：1.曾淑明居民房占用林地勘验图\n2. 曾淑明居民房占用林地小班一览表\n勘验人（签字）：\n重庆市涪陵区林业规划和资源监测中心\n2022年1月12日"
runFirst = paragraphFirst.add_run(strFirst)
runFirst.font.name='方正仿宋_GBK'
runFirst.font.element.rPr.rFonts.set(qn('w:eastAsia'), '方正仿宋_GBK')
runFirst.font.size = Pt(16)
paragraph_formatFirst = paragraphFirst.paragraph_format
paragraph_formatFirst.line_spacing = Pt(28)  # 固定值，30磅
paragraphFirst.paragraph_format.first_line_indent = Cm(1.1)

doc.save('new.docx')
"""
添加段落的时候，赋值给一个变量，方便我们后面进行格式调整；
"""